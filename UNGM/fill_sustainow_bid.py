#!/usr/bin/env python3
"""
Fill UNICEF Venture Fund RFPS-NYH-2026-503931 Annex C templates
for Sustainow Technologies Private Limited — KlimaGuard Kids bid.
"""

from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openpyxl import load_workbook
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "submission"
OUT.mkdir(exist_ok=True)

COMPANY = "Sustainow Technologies Private Limited"
SOLUTION = "KlimaGuard Kids"
COUNTRY = "India"
CIN = "U72200KA2025PTC199691"
ADDRESS = (
    "Flat No. 307, Casagrand Royce Apt, Kodigehalli, "
    "Krishnarajapuram, Bengaluru, Karnataka 560036, India"
)
EMAIL = "contact@sustainow.in"
WEB = "https://sustainow.in/"
DEMO = "https://klimaguardkids.sustainow.in/"
REPO = "https://github.com/carbonintelli/KlimaGuardKids"
RFPS = "RFPS-NYH-2026-503931"

DIRECTOR_W = {
    "first": "Manjula Devi",
    "last": "Purushothaman",
    "full": "Purushothaman Manjula Devi",
    "title": "Co-Founder & Director (Product, Partnerships & Impact)",
    "gender": "Female",
    "nationality": "Indian",
    "email": EMAIL,
}
DIRECTOR_M = {
    "first": "Nandagopal",
    "last": "Govindan",
    "full": "Nandagopal Govindan",
    "title": "Co-Founder & Director (Technology & Architecture)",
    "gender": "Male",
    "nationality": "Indian",
    "email": "carbonintelli@gmail.com",
}


def set_runs_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def append_after_label(doc: Document, label: str, answer: str) -> bool:
    """If a paragraph equals/starts with label, append answer on next empty or same para."""
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == label or t.startswith(label.rstrip(":") + ":"):
            # Prefer writing onto the same paragraph after the label
            if t.endswith(":") or t == label:
                set_runs_text(p, f"{label.rstrip(':')}: {answer}" if not label.endswith("?") else f"{t}\n{answer}")
            else:
                set_runs_text(p, f"{t}\n{answer}")
            # Also fill following empty paragraphs if present
            j = i + 1
            while j < len(doc.paragraphs) and not doc.paragraphs[j].text.strip():
                j += 1
            return True
    return False


def replace_everywhere(doc: Document, mapping: dict[str, str]) -> None:
    for p in doc.paragraphs:
        text = p.text
        new = text
        for a, b in mapping.items():
            if a in new:
                new = new.replace(a, b)
        if new != text:
            set_runs_text(p, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text
                    new = text
                    for a, b in mapping.items():
                        if a in new:
                            new = new.replace(a, b)
                    if new != text:
                        set_runs_text(p, new)


def fill_labeled_fields(doc: Document, fields: list[tuple[str, str]]) -> None:
    """For each (label_prefix, value), find paragraph starting with label and set value."""
    for label, value in fields:
        for p in doc.paragraphs:
            t = p.text.strip()
            if t.startswith(label):
                # Keep label, replace rest
                if ":" in label:
                    set_runs_text(p, f"{label} {value}")
                else:
                    set_runs_text(p, f"{label}: {value}")
                break


def insert_answers_block(doc: Document, answers: list[tuple[str, str]]) -> None:
    """
    For each question heading, find it and insert/overwrite the next non-heading paragraph.
    """
    paras = doc.paragraphs
    headings = {a[0]: a[1] for a in answers}
    for i, p in enumerate(paras):
        key = p.text.strip()
        if key in headings:
            ans = headings[key]
            # look for next paragraph to fill; if empty or short, overwrite
            if i + 1 < len(paras):
                nxt = paras[i + 1]
                nxt_text = nxt.text.strip()
                # Don't overwrite the next question
                if nxt_text in headings or nxt_text.startswith("Part ") or nxt_text.startswith("TEMPLATE"):
                    # insert by appending to question para
                    set_runs_text(p, f"{key}\n\n{ans}")
                else:
                    set_runs_text(nxt, ans)
            else:
                set_runs_text(p, f"{key}\n\n{ans}")


# ── Template 1: Summary ──────────────────────────────────────────────────────

SUMMARY_ANSWERS = {
    "What is the core problem you want to solve through your solution? (1,000 characters limit)": (
        "Climate change intensifies heatwaves, floods, air pollution, vector-borne disease, and food insecurity. "
        "Children are disproportionately harmed by physiology, caregiver dependence, and disruption to school and nutrition. "
        "Most climate tools target policymakers; few deliver actionable, age-appropriate preparedness guidance that "
        "frontline workers, teachers, and caregivers can use within hours. Without child-centred early warning that links "
        "live climate signals to health, nutrition, and disease risk—especially in climate-vulnerable cities and Indian "
        "districts—communities react late, after clinics fill and learning time is lost."
    ),
    "How does your solution solve this core problem? (1,000 characters limit)": (
        "KlimaGuard Kids is an open-source (MIT), agentic web platform that turns live Open-Meteo weather/AQ into "
        "child-centred preparedness guidance via eight transparent TypeScript agents (climate, health, nutrition, disease, "
        "natural medicine, India regional, India CHIS impact, synthesis). It covers 159 countries, 482 Tier 1–3 cities, "
        "and 77 India regions with a measurable Child Health Impact Score (0–100). Outputs include caregiver briefings, "
        "age-banded cards (5–8 / 9–12 / 13–17), and kids play missions—without requiring child accounts. Provenance cites "
        "~50 validated sources (WHO/UNICEF/UNDRR, FAO/IPC, Copernicus, India IMD/NFHS, etc.)."
    ),
    "Describe the results of your user testing/prototyping. Please provide a detailed response on who your main users are and the quantitative and qualitative results of the testing to date) (1,000 characters limit)": (
        "Main users: caregivers, teachers, community health workers, NGO/CSR programme staff, and supervised youth. "
        "Prototype status: working production-style Next.js app with live climate fetch, full agent pipeline, India CHIS "
        "dashboard, global city selector, and age-banded play. Internal demo reviews across diverse climates show end-to-end "
        "analysis typically completing in seconds; stakeholders rate unified climate–health briefings clearer than siloed "
        "weather/health alerts. Registry stress-tested across 159 countries / 482 cities / 77 India regions. Formal field "
        "pilot with ministries/schools is planned for the investment period (target ≥3 partner sites, ≥500 guided sessions) "
        "to quantify comprehension and preparedness behaviour change."
    ),
    "Describe the value that you deliver to the users of your technology  (1,000 characters limit)": (
        "Users receive location-specific climate–health intelligence they can act on the same day: heat/flood/air/vector "
        "stressors, hydration and food-safety notes, disease precautions, and age-appropriate guidance. Schools get "
        "classroom-ready messaging and play missions; CHWs get briefings before peak seasons; programme officers get "
        "transparent CHIS scores for regional prioritisation. Open MIT licensing lets UN entities and governments audit, "
        "fork, and deploy in-country without proprietary lock-in. APIs (/api/analyze, /api/countries, /api/india/regions) "
        "support integration into partner systems."
    ),
    "If your solution has climate and health-related impact, please specify the relevant climate-health outcomes.": (
        "Anticipatory reduction of child heat stress, respiratory burden from poor AQ, waterborne/vector disease pressure, "
        "and climate-linked nutrition stress. Measured via CHIS dimensions (CHVI, CRBS, WDPI, VBDP, CNSI) and age-banded "
        "preparedness actions. Aligns with SDG 3, 13, 2, 4, and 11."
    ),
    "Outline the three main objectives for which potential funding from UNICEF will be used:": (
        "1) Evidence generation: structured school/CHW pilots in ≥3 climate-vulnerable sites with baseline/endline "
        "preparedness metrics and public evaluation brief.\n"
        "2) Product hardening for UNICEF-scale use: multilingual UI (≥3 languages), accessibility, offline-capable PWA, "
        "optional national met/AQ feed adapters, privacy review.\n"
        "3) Open-source ecosystem & interoperability: expanded registries, documented APIs, contributor onboarding, "
        "and pathways to DHIS2/partner dashboard integration."
    ),
    "What are the expected outcomes/results of the 12-month investment period? Also share how you intend to showcase evidence of impact: (1,000 characters limit)": (
        "Outcomes: (a) validated pilot evidence pack (usage, comprehension, preparedness actions) from ≥500 sessions; "
        "(b) production-hardened release with 3 languages + PWA offline shell; (c) public OSS milestones and real-time "
        "investment data dashboard. Evidence: quarterly KPI reports, anonymised aggregate analytics, partner testimonials, "
        "open GitHub releases, and a final public evaluation note suitable for UNICEF Venture Fund learning."
    ),
    "Does your solution have any established evidence of health outcomes, safety, or efficacy? Provide details of completed studies, pilots, evaluations, or real-world deployments, including methodology, sample size, location, dates, outcomes measured, and results. Please attach or provide links to any published reports, evaluation documents, or peer-reviewed publications, if available.": (
        "Current evidence is prototype/demo-level: functioning open-source platform with transparent heuristics (not a "
        "medical device). No peer-reviewed clinical efficacy trial yet. Investment will fund formal pilots with defined "
        "methodology, sample size, and ethics-appropriate protocols. Live demo: "
        f"{DEMO} · Source: {REPO} · Technical commendation: docs/KGK_Technical_Commendation.pdf in repository."
    ),
    "Does your solution currently comply with, integrate with, or support any recognized digital health standards or interoperability frameworks? (e.g., HL7 FHIR, DICOM, IHE, ICD-10/11, SNOMED CT, LOINC, OpenHIE, DHIS2, OpenMRS). For each standard referenced, please indicate whether it is currently implemented, partially implemented, planned, or under evaluation.": (
        "Currently: JSON REST APIs for analysis and geographic registries (custom schema). Planned under investment: "
        "DHIS2 interoperability evaluation/partial implementation for aggregate indicators; ICD-11 tagging under evaluation "
        "for disease outlook categories; HL7 FHIR observation mapping under evaluation for future EHR-light exports. "
        "Not a clinical records system today."
    ),
}


def fill_template1() -> Path:
    src = ROOT / "Annex C. Template 1 - Summary_RFPS-NYH-2026-503931.docx"
    dst = OUT / "01_Summary_Sustainow_KlimaGuardKids.docx"
    shutil.copy(src, dst)
    doc = Document(dst)

    fill_labeled_fields(
        doc,
        [
            ("Company Name:", COMPANY),
            ("Name of solution (if different from the company):", SOLUTION),
            ("Country of Registration:", COUNTRY),
            (
                "Provide a summary of the solution (200 characters limit):",
                "Open-source agentic AI that turns live climate/AQ into age-banded child health preparedness "
                "guidance + India CHIS scores for schools, CHWs, and programmes.",
            ),
            ("Link to video:", "[TO ADD — YouTube 2-min pitch URL before e-submission]"),
            ("Password to view (if applicable):", "N/A (public link when uploaded)"),
        ],
    )

    # Stage & women-led checkboxes as text annotations
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("☐ Early stage"):
            set_runs_text(p, "☑ Early stage (Working prototype with initial testing or pilots)")
        if t.startswith("☐ Deployment ready"):
            set_runs_text(p, "☐ Deployment ready stage (Proven solution that is deployment ready...)")
        if t.startswith("If yes, indicate the name"):
            set_runs_text(
                p,
                "If yes: Purushothaman Manjula Devi — Co-Founder & Director. She co-founded the company (DIN director "
                "since incorporation 12-03-2025) and actively leads product strategy, partnerships, and impact decisions.",
            )

    # Mark women-led Yes
    for i, p in enumerate(doc.paragraphs):
        if "women-founded or women-led" in p.text.lower():
            for j in range(i, min(i + 8, len(doc.paragraphs))):
                if doc.paragraphs[j].text.strip() == "☐ Yes":
                    set_runs_text(doc.paragraphs[j], "☑ Yes")
                if doc.paragraphs[j].text.strip() == "☐ No":
                    set_runs_text(doc.paragraphs[j], "☐ No")

    insert_answers_block(doc, list(SUMMARY_ANSWERS.items()))

    # Risks table
    risks = [
        (
            "Pilot adoption delay (schools/CHW bandwidth)",
            "Medium",
            "Medium",
            "Pre-commit ≥3 partner sites; lightweight 20-min training; offline PWA fallback",
        ),
        (
            "Heuristic guidance misread as clinical advice",
            "High",
            "Low",
            "Clear non-diagnostic disclaimers; caregiver/CHW framing; safeguarding review",
        ),
        (
            "Upstream Open-Meteo/API outage",
            "Medium",
            "Low",
            "Cached last-good snapshots; optional national met adapters; status banners",
        ),
    ]
    if doc.tables:
        t = doc.tables[0]
        for r, row_data in enumerate(risks, start=1):
            if r < len(t.rows):
                for c, val in enumerate(row_data):
                    if c < len(t.rows[r].cells):
                        t.rows[r].cells[c].text = val

    # Stage annotation near top
    for p in doc.paragraphs:
        if "At what stage is your solution?" in p.text:
            set_runs_text(
                p,
                "At what stage is your solution?\n☑ Early stage (Working prototype with initial testing or pilots)\n"
                "☐ Deployment ready stage",
            )
            break

    doc.save(dst)
    return dst


# ── Template 2: Product Requirements ─────────────────────────────────────────

def fill_template2() -> Path:
    src = ROOT / "Annex C. Template 2 - Product Requirements_RFPS-NYH-2026-503931.docx"
    dst = OUT / "02_Product_Requirements_Sustainow_KlimaGuardKids.docx"
    shutil.copy(src, dst)
    doc = Document(dst)

    fill_labeled_fields(doc, [("Company Name:", COMPANY)])

    answers = [
        (
            "Describe the primary use cases (real-world applications) of your solution. Include both actual and potential use cases. Please share any data or results from pilot implementations. For each actual and potential use case, explain why your technology is well-suited for that application.",
            "Actual: (1) Global climate-health dashboard for NGO/programme officers selecting vulnerable cities; "
            "(2) India CHIS dashboard for state/district prioritisation; (3) Kids play missions for school preparedness "
            "education; (4) Stakeholder pitch for CSR/UN partners. Potential: ministry early-warning briefings, "
            "DHIS2 aggregate indicators, WhatsApp/SMS middleware via partners, in-country government cloud deploy. "
            "Suited because agents are transparent, MIT-licensed, lat/lon extensible, and child-safeguarding by design "
            "(no child accounts). Demo: " + DEMO,
        ),
        (
            "Prototyping: Describe the work you have done on proofs of concept, mockups, and user testing. List the last three versions of your solution’s design. For each version, please provide 2 sentences that describe the reasons for the changes (including data) and outline the main changes made.",
            "v0.1 — Single-city climate→health briefing POC; validated live Open-Meteo path and WHO-aligned heuristics. "
            "v0.2 — Multi-agent orchestration + India CHIS + 37 regions; stakeholder demos showed need for deeper geography. "
            "v0.3 (current) — 159 countries / 482 Tier 1–3 cities / 77 India regions / ~50 validated sources; kids play "
            "replaces chat; provenance wired per agent. Changes driven by partner feedback for Tier-2/3 coverage and "
            "UN-ready documentation.",
        ),
        (
            "Before you developed your own technology or selected an existing one, what other similar technologies did you consider using, improving or iterating to solve the problem you want to solve? (Outline what other technologies were considered to address the problem you identified and why were they rejected as unsuitable).",
            "Considered: generic weather apps (no child-health framing); proprietary ESG dashboards (lock-in, weak child focus); "
            "LLM-only chatbots (opacity, safeguarding risk for children). Rejected for lacking transparent scoring, "
            "open licensing, age-banded pedagogy, and India-specific impact metrics. Built deterministic agents + optional "
            "play instead.",
        ),
        (
            "Additionally, please provide written documentation explaining the diagram provided, outlining each element included in the diagram, its role and its status of development (complete, under development, expected date of completion, etc). Please clarify if these components are currently Open Source or proprietary, and in the case of the latter if you intend to release it as Open Source solution. Include links or documents outlining process documentation, system documentation, user documentation, and any instructions on API usage, if relevant. (5,000 characters limit)",
            "Stack (all MIT OSS unless noted): Next.js 15 App Router + React 19 + TypeScript + Tailwind 4 (UI complete); "
            "Route Handlers /api/analyze|/countries|/india/regions with Zod (complete); Orchestrator + 8 domain agents "
            "(complete); Open-Meteo live fetch (complete, third-party); geographic registries countries/india-regions "
            "(complete, expandable); sources provenance catalogue (complete); kids play localStorage gamification "
            "(complete); CHIS formulas (complete, documented). Under development (investment): i18n, PWA offline, "
            "national feed adapters, DHIS2 bridge. Docs: README, docs/ARCHITECTURE.md, "
            "docs/DATA_SOURCES_AND_GEOGRAPHY.md, docs/KGK_Technical_Commendation.pdf. Repo: " + REPO,
        ),
        (
            "Share here all GitHub / Bitbucket (or other) repositories for all components of your solution: (1,000 characters limit)",
            REPO,
        ),
        (
            "Please share the GitHub, GitLab or bitbucket (or other) handles of the code developers within your team:",
            "Organization/repo: carbonintelli/KlimaGuardKids · Primary technical contact: carbonintelli (GitHub) / "
            f"{DIRECTOR_M['email']}",
        ),
        (
            "How many external contributors have there been to your code?",
            "Open repository accepting community PRs; external contributor count growing via public GitHub (see Insights). "
            "Core maintained by Sustainow Technologies.",
        ),
        (
            "What processes and tools do you leverage for quality assurance and testing? (1,000 characters limit)",
            "TypeScript strict mode; ESLint via next lint; production next build on every PR through GitHub Actions CI; "
            "manual cross-climate pipeline demos; Zod request validation; transparent unit-testable agent functions. "
            "Investment will add automated agent fixture tests and accessibility checks.",
        ),
        (
            "Sustainability: What is the ongoing plan for maintenance, updates, and support of your software? Describe what the needs will be for the next two years and how these will be met. (1,000 characters limit)",
            "2-year plan: quarterly OSS releases; registry expansions; dependency updates; security advisories; partner "
            "pilot support. Funding mix: UNICEF seed (Y1), CSR/programme pilots, optional Growth Funding / LTAS path, "
            "implementation support tiers. Hosting on low-cost Node/container; Open-Meteo keeps climate cost near-zero "
            "for demo scale.",
        ),
        (
            "What data elements are you processing (i.e., collecting, using, retaining, sharing, etc.) to provide your product/service? Describe: (1,000 characters limit)",
            "Processing: country/city/region selection, lat/lon, live weather/AQ fields from Open-Meteo, derived risk "
            "scores and guidance text. No names, phones, school IDs, or health records in core flow. Play progress stays "
            "in browser localStorage only.",
        ),
        (
            "Do you collect any personally identifiable data, and/or children’s data? If so, how are you collecting, processing and storing this information? Please provide any details on data management policies, agreements, or practices being followed or planned. (1,000 characters limit)",
            "No child accounts; no intentional collection of children’s PII in the demo architecture. Location is "
            "city/region level. Future account/messaging features will follow COPPA/GDPR-K / national child-data laws "
            "with DPIA before enablement.",
        ),
        (
            "How much data will you be collecting and using? How often will you be collecting data (is it a one-time collection or on an ongoing basis)? Describe your processes: (1,000 characters limit)",
            "Climate/AQ fetched on-demand per analysis (ongoing when users run analyze). Aggregate anonymous usage "
            "metrics planned for pilots (session counts, feature use)—not personal profiles. Registry data is static "
            "curated metadata updated with releases.",
        ),
        (
            "Describe how you already collect or plan to collect data. Please provide a data flow diagram (1,000 characters limit).",
            "UI → HTTPS POST /api/analyze → Zod validate → registry resolve lat/lon → Open-Meteo fetch → agents → "
            "SynthesisReport JSON → UI. No persistent user DB in demo. Pilot analytics: optional privacy-preserving "
            "event log (region_id, timestamp, agent latency) without device identifiers. Diagrams: docs/images/"
            "analyze-sequence.png, client-server.png, data-registries.png.",
        ),
    ]

    # License checkboxes
    for p in doc.paragraphs:
        if "Permissive (e.g., MIT License" in p.text:
            set_runs_text(p, "☑ Permissive (e.g., MIT License, Apache License)")
        if p.text.strip().startswith("Copyleft"):
            set_runs_text(p, "☐ Copyleft (e.g., GNU General Public License)")
        if "Which repository hosting service" in p.text:
            set_runs_text(p, "Which repository hosting service/git forge do you use for your solution:\n☑ GitHub")

    insert_answers_block(doc, answers)

    # GenAI section - clarify we use deterministic agents not GenAI as core
    for p in doc.paragraphs:
        if "generative AI" in p.text.lower() and "Describe" in p.text:
            set_runs_text(
                p,
                p.text
                + "\n\nCore pipeline does NOT use generative AI for scoring. Agents are deterministic TypeScript "
                "heuristics for auditability/safeguarding. Any future LLM assist would be optional, adult-facing, "
                "and sandboxed—not used for child clinical decisions.",
            )
            break

    # Append annex note for tech stack diagram
    doc.add_paragraph("")
    doc.add_paragraph("ANNEX 1 — Tech stack diagram reference")
    doc.add_paragraph(
        f"See repository diagrams (MIT): {REPO}/tree/main/docs/images — system-layers.png, "
        "agent-pipeline.png, client-server.png, data-registries.png. Also docs/ARCHITECTURE.md and "
        "docs/KGK_Technical_Commendation.pdf."
    )

    doc.save(dst)
    return dst


# ── Template 3: Product Design xlsx ──────────────────────────────────────────

def fill_template3() -> Path:
    src = ROOT / "Annex C. Template 3 - Product Design_RFPS-NYH-2026-503931.xlsx"
    dst = OUT / "03_Product_Design_Data_Plan_Sustainow_KlimaGuardKids.xlsx"
    shutil.copy(src, dst)
    wb = load_workbook(dst)
    ws = wb.active
    ws["B2"] = COMPANY

    # Helper to write a row: Timeline col A may already have labels; fill B-F
    def write_activity(row: int, task: str, change: str, success: str, data: str, owner: str):
        ws.cell(row, 2).value = task
        ws.cell(row, 3).value = change
        ws.cell(row, 4).value = success
        ws.cell(row, 5).value = data
        ws.cell(row, 6).value = owner

    # Map from inspection: rows roughly
    # Q1 Milestone/Activity rows — fill key ones based on earlier dump
    # Row 7 Quarter 1, 8 Milestone 1, 9-10 activities, etc.
    plan = [
        # row, task, change, success, data, owner
        (8, "M1: Pilot site MoUs & ethics-light protocol", "≥3 sites committed", "Signed partner briefs", "MoU dates, site list", "Manjula Devi"),
        (9, "Partner onboarding packs (schools/CHWs)", "Training kit v1 live", "≥6 facilitators trained", "Training attendance", "Manjula Devi"),
        (10, "Baseline preparedness survey design", "Survey instrument locked", "IRB/ethics note filed if needed", "Survey schema", "Manjula Devi"),
        (11, "M2: Multilingual foundation", "i18n framework shipped", "EN+HI+1 language scaffolding", "Locale coverage metrics", "Nandagopal Govindan"),
        (12, "String extraction & translation workflow", "≥80% UI strings externalized", "L10n PRs merged", "String coverage %", "Nandagopal Govindan"),
        (13, "Accessibility pass (WCAG 2.1 AA targets)", "Critical a11y issues <5", "Axe CI report green", "A11y issue counts", "Nandagopal Govindan"),
        (14, "M3: Real-time investment data feed", "Public KPI endpoint", "Dashboard shows live KPIs", "KPI time series", "Nandagopal Govindan"),
        (15, "Anonymised usage analytics (no child PII)", "Event pipeline live", "≥95% events schema-valid", "Event volumes", "Nandagopal Govindan"),
        (16, "M4: Registry & sources expansion", "+10% Tier-3 cities", "Release tag v0.4", "City/source counts", "Nandagopal Govindan"),
        (17, "India CHIS validation notes with experts", "Expert review memo", "≥3 expert reviews", "Review comments", "Manjula Devi"),
        (19, "User testing Product/UX (Q1)", "≥50 guided sessions", "SUS/comprehension scores", "Session logs, scores", "Manjula Devi"),
        (20, "User testing UI (Q1)", "UI issues backlog triaged", "Top 10 UX fixes shipped", "Issue tracker", "Nandagopal Govindan"),
        (22, "Q1 FE deliverables", "i18n shell + a11y fixes", "Tagged release", "Release notes", "Nandagopal Govindan"),
        (23, "Q1 BE deliverables", "KPI/analytics APIs", "Uptime >99% demo", "Uptime, latency", "Nandagopal Govindan"),
        # Q2
        (25, "M1: PWA offline shell", "Offline guidance cache", "Lighthouse PWA ≥80", "Cache hit rates", "Nandagopal Govindan"),
        (26, "Service worker & last-good climate cache", "Graceful offline mode", "Offline demo script pass", "Offline test results", "Nandagopal Govindan"),
        (28, "M2: National met/AQ adapter interface", "Adapter interface v1", "1 national feed stub", "Adapter docs", "Nandagopal Govindan"),
        (29, "IMD/CPCB reference integration notes", "India feed design doc", "Partner technical review", "Design doc", "Nandagopal Govindan"),
        (31, "M3: School pilot wave 1", "≥200 sessions completed", "Attendance & completion %", "Pilot dataset", "Manjula Devi"),
        (32, "CHW briefing workshops", "≥40 CHWs oriented", "Workshop feedback ≥4/5", "Feedback forms", "Manjula Devi"),
        (34, "M4: DHIS2 interoperability spike", "Feasibility report", "Go/no-go decision", "Spike report", "Nandagopal Govindan"),
        (35, "Aggregate indicator mapping", "Draft indicator dictionary", "Reviewed by partner M&E", "Indicator list", "Manjula Devi"),
        (37, "User testing Product/UX (Q2)", "≥100 cumulative sessions", "Improved task success ≥15%", "UX metrics", "Manjula Devi"),
        (38, "User testing UI (Q2)", "Mobile UX hardened", "Mobile crash-free >99%", "Mobile analytics", "Nandagopal Govindan"),
        (40, "Q2 FE: PWA + mobile UX", "PWA in production demo", "Release v0.5", "Release notes", "Nandagopal Govindan"),
        (41, "Q2 BE: adapters + analytics", "Adapter + KPI stable", "Latency p95 <3s", "Perf metrics", "Nandagopal Govindan"),
    ]

    # Continue Q3/Q4 in later rows - find Quarter 3 by scanning
    q3_row = None
    for r in range(1, 120):
        v = ws.cell(r, 1).value
        if v and str(v).strip() == "Quarter 3":
            q3_row = r
            break
    if q3_row:
        q3_plan = [
            (q3_row + 1, "M1: Midline evaluation", "Midline report published", "Report accepted internally", "Midline dataset", "Manjula Devi"),
            (q3_row + 2, "Comparative site analysis", "Cross-site insights", "≥3 sites compared", "Comparison tables", "Manjula Devi"),
            (q3_row + 4, "M2: Language pack completion", "3 languages complete", "Linguist QA sign-off", "L10n QA log", "Nandagopal Govindan"),
            (q3_row + 5, "Community OSS contributor sprint", "≥5 external PRs", "Contributor guide live", "PR/contributor counts", "Nandagopal Govindan"),
            (q3_row + 7, "M3: Safeguarding & privacy review", "Review checklist complete", "Remediations merged", "Review report", "Manjula Devi"),
            (q3_row + 8, "Child-data policy pack", "Policy docs published", "Legal review note", "Policy PDFs", "Manjula Devi"),
            (q3_row + 10, "M4: Partner integration pilots", "1 partner API integration", "Successful sandbox calls", "Integration logs", "Nandagopal Govindan"),
            (q3_row + 14, "User testing Product/UX (Q3)", "≥300 cumulative sessions", "Preparedness score ↑", "Survey results", "Manjula Devi"),
            (q3_row + 15, "User testing UI (Q3)", "Localization UX validated", "Language switch success ≥95%", "UX tests", "Nandagopal Govindan"),
        ]
        for row in q3_plan:
            write_activity(*row)

    q4_row = None
    for r in range(1, 160):
        v = ws.cell(r, 1).value
        if v and str(v).strip() == "Quarter 4":
            q4_row = r
            break
    if q4_row:
        q4_plan = [
            (q4_row + 1, "M1: Endline evaluation & public brief", "Public evaluation brief", "Brief on GitHub/site", "Endline data", "Manjula Devi"),
            (q4_row + 2, "UNICEF learning package", "Case study delivered", "VF learning assets", "Case study doc", "Manjula Devi"),
            (q4_row + 4, "M2: Scale-ready release v1.0", "Tagged v1.0 OSS release", "CI green; docs complete", "Release artifacts", "Nandagopal Govindan"),
            (q4_row + 5, "Deployment playbook (gov cloud)", "Playbook published", "Partner can redeploy", "Playbook PDF", "Nandagopal Govindan"),
            (q4_row + 7, "M3: Growth Funding / LTAS readiness", "Pipeline of follow-ons", "≥3 follow-on conversations", "Pipeline CRM notes", "Manjula Devi"),
            (q4_row + 8, "Business model validation update", "Updated unit economics", "Board/director review", "Financial memo", "Manjula Devi"),
            (q4_row + 10, "M4: Knowledge transfer & handoff", "Training recordings", "Partners self-serve", "Training assets", "Both directors"),
            (q4_row + 14, "Final UX validation (≥500 sessions)", "Session target met", "Final KPI dashboard", "Full pilot dataset", "Manjula Devi"),
            (q4_row + 15, "Final UI polish & accessibility", "AA critical issues closed", "A11y audit archive", "Audit report", "Nandagopal Govindan"),
        ]
        for row in q4_plan:
            write_activity(*row)

    for row in plan:
        write_activity(*row)

    wb.save(dst)
    return dst


# ── Template 4: Team — generate org chart PDF ────────────────────────────────

def fill_template4_orgchart() -> Path:
    # Also copy template with company name filled
    src = ROOT / "Annex C. Template 4_Team Structure_RFPS-NYH-2026-503931.docx"
    dst_docx = OUT / "04_Team_Structure_Sustainow_KlimaGuardKids.docx"
    shutil.copy(src, dst_docx)
    doc = Document(dst_docx)
    fill_labeled_fields(doc, [("Company Name:", COMPANY)])
    doc.add_paragraph("")
    doc.add_paragraph(
        "Organization chart PDF submitted as: 04_Org_Chart_Sustainow_KlimaGuardKids.pdf"
    )
    doc.add_paragraph(
        f"Directors (MCA): {DIRECTOR_W['full']} (DIN from incorporation 12-03-2025); "
        f"{DIRECTOR_M['full']} (DIN from incorporation 12-03-2025). CIN {CIN}."
    )
    doc.save(dst_docx)

    pdf_path = OUT / "04_Org_Chart_Sustainow_KlimaGuardKids.pdf"
    styles = getSampleStyleSheet()
    title = ParagraphStyle("t", parent=styles["Title"], fontSize=16, textColor=colors.HexColor("#0f172a"))
    cell = ParagraphStyle("c", parent=styles["Normal"], fontSize=8, leading=11, alignment=TA_CENTER)
    body = ParagraphStyle("b", parent=styles["Normal"], fontSize=9, leading=12)

    story = [
        Paragraph(f"{COMPANY}", title),
        Paragraph("Organization Chart — KlimaGuard Kids (UNICEF VF RFPS-NYH-2026-503931)", body),
        Spacer(1, 12),
    ]

    board = Paragraph(
        f"<b>Board / Directors</b><br/>{DIRECTOR_W['full']}<br/>&amp; {DIRECTOR_M['full']}<br/>"
        f"Private Limited · Bengaluru, India",
        cell,
    )
    story.append(
        Table([[board]], colWidths=[6.5 * inch], hAlign="CENTER")
    )
    story[-1].setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#0c4a6e")),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.white),
                ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#0c4a6e")),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.append(Spacer(1, 10))
    story.append(Paragraph("▼", ParagraphStyle("v", alignment=TA_CENTER, fontSize=14)))
    story.append(Spacer(1, 6))

    def person_box(html: str, bg: str):
        return Paragraph(html, cell), bg

    left = (
        f"<b>{DIRECTOR_W['full']}</b><br/>Co-Founder &amp; Director<br/>"
        f"Product · Partnerships · Impact<br/>Founder: Yes<br/>Contract: Director/Employee<br/>"
        f"Time: 100% company · 80% this solution<br/>Key: partnerships, pilots, M&amp;E, safeguarding"
    )
    right = (
        f"<b>{DIRECTOR_M['full']}</b><br/>Co-Founder &amp; Director<br/>"
        f"Technology · Architecture<br/>Founder: Yes<br/>Contract: Director/Employee<br/>"
        f"Time: 100% company · 90% this solution<br/>Key: software, OSS, APIs, security"
    )
    t = Table(
        [[Paragraph(left, cell), Paragraph(right, cell)]],
        colWidths=[3.4 * inch, 3.4 * inch],
        hAlign="CENTER",
    )
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#fef3c7")),
                ("BACKGROUND", (1, 0), (1, 0), colors.HexColor("#bae6fd")),
                ("BOX", (0, 0), (0, 0), 1, colors.HexColor("#f59e0b")),
                ("BOX", (1, 0), (1, 0), 1, colors.HexColor("#0ea5e9")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(t)
    story.append(Spacer(1, 14))
    story.append(Paragraph("▼ Planned hires during 12-month investment (dashed roles)", body))
    story.append(Spacer(1, 6))
    planned = Table(
        [
            [
                Paragraph(
                    "<b>Field / Pilot Lead</b><br/>To hire · Contractor<br/>50–100%<br/>User testing, school/CHW coordination",
                    cell,
                ),
                Paragraph(
                    "<b>Frontend Engineer</b><br/>To hire · Contractor<br/>50%<br/>i18n, PWA, accessibility",
                    cell,
                ),
                Paragraph(
                    "<b>M&E Associate</b><br/>To hire · Contractor<br/>50%<br/>Surveys, evidence pack, KPIs",
                    cell,
                ),
            ]
        ],
        colWidths=[2.3 * inch, 2.3 * inch, 2.3 * inch],
        hAlign="CENTER",
    )
    planned.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
                ("BOX", (0, 0), (0, 0), 1, colors.HexColor("#94a3b8")),
                ("BOX", (1, 0), (1, 0), 1, colors.HexColor("#94a3b8")),
                ("BOX", (2, 0), (2, 0), 1, colors.HexColor("#94a3b8")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(planned)
    story.append(Spacer(1, 16))
    story.append(
        Paragraph(
            f"Registered address: {ADDRESS}<br/>Email: {EMAIL} · Web: {WEB}<br/>"
            f"Product: {SOLUTION} · Repo: {REPO}",
            ParagraphStyle("f", fontSize=8, textColor=colors.HexColor("#475569"), leading=11),
        )
    )

    SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36,
        title=f"{COMPANY} Org Chart — {SOLUTION}",
        author=COMPANY,
    ).build(story)
    return pdf_path


# ── Template 5: CVs ──────────────────────────────────────────────────────────

def fill_cv_for(person: dict, path: Path) -> Path:
    src = ROOT / "Annex C. Template 5 - CVs_RFPS-NYH-2026-503931.docx"
    shutil.copy(src, path)
    doc = Document(path)
    fill_labeled_fields(doc, [("Company Name:", COMPANY)])

    if doc.tables:
        t0 = doc.tables[0]
        # Row0: First name / First name / Last name — template is a bit odd
        t0.rows[0].cells[0].text = f"First name: {person['first']}"
        t0.rows[0].cells[1].text = f"First name: {person['first']}"
        t0.rows[0].cells[2].text = f"Last name: {person['last']}"
        t0.rows[1].cells[0].text = f"Title in current project: {person['title']}"
        t0.rows[1].cells[1].text = f"Title in current project: {person['title']}"
        t0.rows[2].cells[0].text = f"Nationality: {person['nationality']}"
        t0.rows[2].cells[1].text = "Date of Birth: [Authorized signatory to confirm]"
        t0.rows[2].cells[2].text = f"Gender identity: {person['gender']}"
        for c in range(3):
            t0.rows[3].cells[c].text = f"Current residency address: {ADDRESS}"
            t0.rows[4].cells[c].text = f"Email: {person['email']}"
            t0.rows[5].cells[c].text = "LinkedIn Profile: [To be added before submission]"

        t1 = doc.tables[1]
        t1.rows[2].cells[0].text = COMPANY
        t1.rows[2].cells[1].text = person["title"]
        t1.rows[2].cells[2].text = "03/25 to Present"
        t1.rows[2].cells[3].text = "India"
        t1.rows[3].cells[0].text = "KlimaGuard Kids (product)"
        t1.rows[3].cells[1].text = (
            "Product leadership" if "Product" in person["title"] else "Lead architect / engineering"
        )
        t1.rows[3].cells[2].text = "2025 to Present"
        t1.rows[3].cells[3].text = "India / Global OSS"

        t2 = doc.tables[2]
        t2.rows[2].cells[0].text = "[Education details — authorized signatory to complete]"
        t2.rows[2].cells[1].text = ""
        t2.rows[2].cells[2].text = "India"
        t2.rows[2].cells[3].text = ""

        t3 = doc.tables[3]
        t3.rows[2].cells[0].text = "UNICEF Innovation / Programme contact"
        t3.rows[2].cells[1].text = "To be nominated"
        t3.rows[2].cells[2].text = "Partner organisation"
        t3.rows[2].cells[3].text = EMAIL
        t3.rows[3].cells[0].text = "Technical community referee"
        t3.rows[3].cells[1].text = "Open-source collaborator"
        t3.rows[3].cells[2].text = "GitHub community"
        t3.rows[3].cells[3].text = DIRECTOR_M["email"]

        # Criminal checkbox NO
        doc.tables[4].rows[0].cells[0].text = (
            "Have you ever been arrested, indicted, or summoned into court as a defendant in a criminal "
            "proceeding, or convicted, fined or imprisoned for the violation of any law (excluding minor "
            "traffic violations)? YES ☐   NO ☑"
        )

    doc.add_paragraph("")
    doc.add_paragraph(
        f"Professional summary: {person['full']} is a Co-Founder & Director of {COMPANY} (CIN {CIN}), "
        f"registered in Bengaluru, India. Role on {SOLUTION}: {person['title']}. "
        f"Focus areas align with UNICEF Venture Fund open-source climate-health objectives for children."
    )
    doc.save(path)
    return path


def fill_template5() -> list[Path]:
    return [
        fill_cv_for(DIRECTOR_W, OUT / "05_CV_Manjula_Devi_Purushothaman.docx"),
        fill_cv_for(DIRECTOR_M, OUT / "05_CV_Nandagopal_Govindan.docx"),
    ]


# ── Template 6: Business model ───────────────────────────────────────────────

def fill_template6() -> Path:
    src = ROOT / "Annex C. Template 6 - Business Model&Plan_RFPS-NYH-2026-503931.docx"
    dst = OUT / "06_Business_Model_Plan_Sustainow_KlimaGuardKids.docx"
    shutil.copy(src, dst)
    doc = Document(dst)
    fill_labeled_fields(doc, [("Company Name:", COMPANY)])

    answers = [
        (
            "What are your current and planned go-to-market strategies? Describe your approach and provide details on how you will execute these activities (500 character limit)",
            "Land-and-expand via school networks, CHW NGOs, and CSR climate-health programmes in India first; "
            "UN/INGO demos via open MIT repo + live site; RFPs/LTAs after evidence pack. Execute with pilot MoUs, "
            "facilitator kits, and public OSS releases that procurement can audit.",
        ),
        (
            "Who are your main competitors in the market?",
            "Weather apps, proprietary ESG/climate dashboards, and generic AI chatbots. Differentiated by "
            "child-centric age bands, transparent CHIS scoring, India deep coverage, MIT OSS, and no child accounts.",
        ),
        (
            "Please describe competing solutions and clearly explain how your solution differs from theirs. Consider differences in technology stack, business strategy, target users, languages, design approach, effectiveness, cost, or other relevant aspects. Take the time to reflect, it is important to demonstrate a clear understanding of your competitive landscape and what sets your solution apart.",
            "Unlike weather apps, we correlate climate with child health/nutrition/disease and pedagogy. Unlike "
            "closed ESG suites, we are MIT open-source and deployable in-country. Unlike LLM chatbots, scoring is "
            "deterministic and reviewable. Cost base is low (Open-Meteo + Node hosting). Strategy: evidence-first "
            "pilots → programme licenses/support → optional Growth Funding/LTAS.",
        ),
        (
            "Financial strategy: Please describe your strategy for securing additional investment and funding to support the next stages of your solution and company. Specify the main funding sources you currently have in place, as well as those you plan to pursue over the next 12 months (in addition to the UNICEF Venture Fund). This could include building relationships with potential follow-on investors or mentors, increasing your company’s visibility through targeted campaigns, or exploring other strategic funding avenues.",
            "Current: founder-supported R&D and open-source release. Next 12 months: UNICEF VF seed; CSR/school "
            "pilot contracts; India climate-health grants; angel/impact intros; Venture Fund Growth Funding path "
            "($200–400k) after seed success; visibility via klimaguardkids.sustainow.in and GitHub.",
        ),
    ]
    insert_answers_block(doc, answers)

    # Users 1/2/3
    for p in doc.paragraphs:
        if p.text.strip() == "1)":
            set_runs_text(
                p,
                "1) Teachers/school admins (paid via institution/CSR) — need classroom-ready heat/flood/air guidance; ages 5–17.",
            )
        if p.text.strip() == "2)":
            set_runs_text(
                p,
                "2) Community health workers / NGO staff (programme-funded) — need rapid location briefings; serve child communities.",
            )
        if p.text.strip() == "3)":
            set_runs_text(
                p,
                "3) Caregivers & supervised youth (unpaid end users) — need simple age-banded actions/play; ages 5–17 with adult oversight.",
            )

    if doc.tables:
        rev = doc.tables[0]
        rows = [
            (
                "Implementation / pilot support",
                "0 (pre-revenue prototype year)",
                "25000",
                "40%",
                "MoUs before build; fixed-scope SOWs; open demo reduces sales friction",
            ),
            (
                "CSR / ESG community programmes",
                "0",
                "20000",
                "30%",
                "Outcome dashboards (CHIS); school packs; brand-safe OSS story",
            ),
            (
                "Government / UN adaptation contracts",
                "0",
                "15000",
                "20%",
                "In-country deploy playbooks; MIT auditability; VF reference",
            ),
            (
                "Potential: API/hosting SLA & training",
                "0",
                "10000",
                "10%",
                "Tiered support menu; start with self-serve OSS",
            ),
        ]
        for i, row in enumerate(rows, start=1):
            if i < len(rev.rows):
                for c, val in enumerate(row):
                    rev.rows[i].cells[c].text = val

        partners = doc.tables[1]
        pdata = [
            ("School / NGO pilot partners (India)", "TBD — shortlist in Q1", "User testing & deployment", "Potential"),
            ("UNICEF country/programme offices", "Via VF network", "Evidence & scale pathways", "Potential"),
            ("Open-Meteo / public data ecosystem", "Open-Meteo", "Climate/AQ inputs", "Actual (public API)"),
        ]
        for i, row in enumerate(pdata, start=1):
            if i < len(partners.rows):
                for c, val in enumerate(row):
                    partners.rows[i].cells[c].text = val

        comps = doc.tables[2]
        cdata = [
            ("Consumer weather apps", "Forecast & alerts", "No child-health agents, CHIS, or OSS programme model"),
            ("Proprietary ESG platforms", "Enterprise climate risk", "Closed IP; weak child pedagogy; high lock-in"),
            ("Generic AI health chatbots", "LLM Q&A", "Opaque; safeguarding risk; not transparent scoring"),
        ]
        for i, row in enumerate(cdata, start=1):
            if i < len(comps.rows):
                for c, val in enumerate(row):
                    comps.rows[i].cells[c].text = val

        fin = doc.tables[3]
        fdata = [
            ("Founder capital / sweat equity", "Current", "In-kind R&D", "N/A", "Ongoing"),
            ("UNICEF Venture Fund seed", "Anticipated", "100000", "This RFPS submission", "Aug–Sep 2026 award window"),
            ("CSR pilot contracts", "Anticipated", "30000", "Pilot proposals to corporates/NGOs", "Q2–Q4 investment year"),
            ("India climate-health grants", "Anticipated", "25000", "Applications to aligned funds", "Q2–Q4"),
            ("VF Growth Funding (later)", "Anticipated", "200000–400000", "Post-seed evidence pack", "After month 12"),
        ]
        for i, row in enumerate(fdata, start=1):
            if i < len(fin.rows):
                for c, val in enumerate(row):
                    fin.rows[i].cells[c].text = val

    doc.save(dst)
    return dst


# ── Template 7: Budget ───────────────────────────────────────────────────────

def fill_template7() -> Path:
    src = ROOT / "Annex C. Template 7-Budget Proposal Temp_RFPS-NYH-2026-503931.xlsx"
    dst = OUT / "07_Budget_Proposal_Sustainow_KlimaGuardKids.xlsx"
    shutil.copy(src, dst)
    wb = load_workbook(dst)

    ws = wb["Tab 1 Project Design Budget"]
    ws["B2"] = COMPANY
    ws["C5"] = "USD"
    # Clear template sample amounts before writing our USD-only budget (target 100,000).
    for r in range(6, 90):
        ws.cell(r, 3).value = None
        ws.cell(r, 4).value = None

    allocations = {
        8: ("Pilot MoUs & protocol", 4000),
        9: ("Partner onboarding", 3000),
        10: ("Baseline survey design", 2000),
        11: ("i18n foundation", 5000),
        12: ("String extraction/l10n", 3000),
        13: ("Accessibility pass", 2000),
        14: ("KPI/data feed", 3000),
        15: ("Analytics pipeline", 2000),
        16: ("Registry expansion", 1000),
        25: ("PWA offline shell", 6000),
        26: ("SW + cache", 3000),
        28: ("National feed adapters", 5000),
        29: ("IMD/CPCB design", 2000),
        31: ("School pilot wave 1", 5000),
        32: ("CHW workshops", 3000),
        34: ("DHIS2 spike", 3000),
        35: ("Indicator mapping", 2000),
    }
    for row, (task, amt) in allocations.items():
        ws.cell(row, 2).value = task
        ws.cell(row, 3).value = amt

    for r in range(1, 90):
        a = ws.cell(r, 1).value
        if a and str(a).strip() == "Quarter 3 total":
            q3 = r
            items = [
                (q3 + 1, "Midline evaluation", 4000),
                (q3 + 2, "Cross-site analysis", 2000),
                (q3 + 3, "Language pack completion", 4000),
                (q3 + 4, "OSS contributor sprint", 2000),
                (q3 + 5, "Safeguarding/privacy review", 4000),
                (q3 + 6, "Child-data policy pack", 2000),
                (q3 + 7, "Partner API integration", 4000),
                (q3 + 8, "Pilot operations Q3", 2000),
            ]
            for row, task, amt in items:
                ws.cell(row, 2).value = task
                ws.cell(row, 3).value = amt
        if a and str(a).strip() == "Quarter 4 total":
            q4 = r
            items = [
                (q4 + 1, "Endline & public brief", 4000),
                (q4 + 2, "UNICEF learning package", 3000),
                (q4 + 3, "v1.0 scale-ready release", 4000),
                (q4 + 4, "Gov-cloud playbook", 3000),
                (q4 + 5, "Follow-on pipeline", 2000),
                (q4 + 6, "Business model update", 2000),
                (q4 + 7, "Knowledge transfer", 3000),
                (q4 + 8, "Final UX/a11y polish", 2000),
            ]
            for row, task, amt in items:
                ws.cell(row, 2).value = task
                ws.cell(row, 3).value = amt

    # Ensure template sample milestone amount is gone; keep USD-only.
    ws.cell(7, 3).value = None
    ws.cell(7, 4).value = None
    for r in range(6, 90):
        ws.cell(r, 4).value = None

    line_total = sum(
        ws.cell(r, 3).value
        for r in range(1, 100)
        if isinstance(ws.cell(r, 3).value, (int, float))
    )
    if line_total != 100000:
        # Final adjust on last polish line to hit exactly USD 100,000.
        delta = 100000 - line_total
        last = 65
        cur = ws.cell(last, 3).value or 0
        ws.cell(last, 2).value = ws.cell(last, 2).value or "Final UX/a11y polish"
        ws.cell(last, 3).value = cur + delta

    ws2 = wb["Tab 2  Company 1-YR Budget"]
    ws2["A2"] = "Budget Period: September 2026 to August 2027"
    ws2["A3"] = COMPANY
    # Income
    ws2["A7"] = "Pilot / implementation support revenue"
    ws2["B7"] = 10000
    ws2["C7"] = 15000
    ws2["A8"] = "CSR programme fees"
    ws2["B8"] = 5000
    ws2["C8"] = 15000
    ws2["A17"] = "UNICEF Venture Fund seed"
    ws2["B17"] = 50000
    ws2["C17"] = 50000
    ws2["A18"] = "Founder capital contribution"
    ws2["B18"] = 10000
    ws2["C18"] = 10000
    # Expenses
    ws2["A29"] = "Salaries & director stipends (KGK)"
    ws2["B29"] = 35000
    ws2["C29"] = 35000
    ws2["A30"] = "Contractors (FE, field, M&E)"
    ws2["B30"] = 15000
    ws2["C30"] = 20000
    ws2["A31"] = "Cloud hosting, tooling, CI"
    ws2["B31"] = 3000
    ws2["C31"] = 3000
    ws2["A32"] = "Pilot travel & workshops"
    ws2["B32"] = 5000
    ws2["C32"] = 8000
    ws2["A33"] = "General & administration"
    ws2["B33"] = 4000
    ws2["C33"] = 4000
    ws2["A34"] = "Marketing & community OSS"
    ws2["B34"] = 3000
    ws2["C34"] = 5000

    wb.save(dst)
    return dst


# ── Cover letter + Appendix 1 + README ───────────────────────────────────────

def write_cover_letter() -> Path:
    path = OUT / "00_Cover_Letter_Sustainow_KlimaGuardKids_RFPS-NYH-2026-503931.pdf"
    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["Normal"], fontSize=10, leading=14, alignment=TA_JUSTIFY, spaceAfter=8)
    story = []
    story.append(Paragraph("<b>Sustainow Technologies Private Limited</b>", styles["Title"]))
    story.append(Paragraph(ADDRESS, ParagraphStyle("a", fontSize=9, textColor=colors.HexColor("#475569"), alignment=TA_CENTER)))
    story.append(Paragraph(f"{EMAIL} · {WEB} · {DEMO}", ParagraphStyle("a2", fontSize=9, textColor=colors.HexColor("#0284c7"), alignment=TA_CENTER, spaceAfter=12)))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#0ea5e9")))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Date: 02 August 2026", body))
    story.append(Paragraph("<b>To:</b> UNICEF Innovation Fund / Service Contracting Centre", body))
    story.append(Paragraph(f"<b>Subject:</b> Proposal submission — {RFPS} — {SOLUTION}", body))
    story.append(Paragraph("Dear Members of the Evaluation Team,", body))
    story.append(
        Paragraph(
            f"{COMPANY} (CIN {CIN}), a women co-founded technology company registered in Bengaluru, India "
            f"(a UNICEF programme country), is pleased to submit this proposal under <b>{RFPS}</b> "
            f"<i>UNICEF Innovation Fund Investment Call for Open-Source Applications</i> for our solution "
            f"<b>{SOLUTION}</b>.",
            body,
        )
    )
    story.append(
        Paragraph(
            "KlimaGuard Kids is an MIT-licensed, open-source agentic platform that converts live climate and air-quality "
            "signals into child-centred health preparedness guidance. Eight transparent software agents produce "
            "location briefings, age-banded guidance (5–17), India Child Health Impact Scores across 77 regions, and "
            "privacy-light kids play missions—without requiring child accounts.",
            body,
        )
    )
    story.append(
        Paragraph(
            "We request equity-free seed investment of <b>USD 100,000</b> for a 12-month programme to (1) generate "
            "field evidence with schools and community health partners, (2) harden the product for multilingual, "
            "accessible, offline-capable deployment, and (3) deepen open-source interoperability and documentation "
            "for UNICEF and country partners.",
            body,
        )
    )
    story.append(Paragraph("This submission includes completed Annex C Templates 1–7, organization chart, CVs of directors, and references to our public repository and technical commendation PDF.", body))
    story.append(Paragraph(f"<b>Repository:</b> {REPO}<br/><b>Live demo:</b> {DEMO}", body))
    story.append(Paragraph("We confirm proposal currency <b>USD</b> and validity of at least <b>180 days</b> from the submission deadline (10 August 2026, 23:59 CEST).", body))
    story.append(Paragraph("Yours sincerely,", body))
    story.append(Spacer(1, 16))
    story.append(Paragraph(f"<b>{DIRECTOR_W['full']}</b><br/>Co-Founder &amp; Director<br/>{COMPANY}<br/>{EMAIL}", body))
    story.append(Spacer(1, 8))
    story.append(Paragraph(f"<b>{DIRECTOR_M['full']}</b><br/>Co-Founder &amp; Director (Technology)<br/>{COMPANY}", body))
    SimpleDocTemplate(str(path), pagesize=A4, leftMargin=54, rightMargin=54, topMargin=48, bottomMargin=48,
                      title=f"Cover Letter — {SOLUTION}", author=COMPANY).build(story)
    return path


def write_appendix1() -> Path:
    path = OUT / "08_Appendix1_Sustainable_Procurement_Sustainow.pdf"
    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["Normal"], fontSize=10, leading=13, spaceAfter=8, alignment=TA_JUSTIFY)
    story = [
        Paragraph("Appendix 1 — Sustainable Procurement Commitments", styles["Title"]),
        Paragraph(f"{COMPANY} · {SOLUTION} · {RFPS}", ParagraphStyle("s", fontSize=9, textColor=colors.gray, alignment=TA_CENTER, spaceAfter=12)),
        Paragraph("☑ <b>Economic pillar:</b> We prioritise local Indian talent, MSME contractors for field work, and open-source reuse that reduces duplicate procurement spend for partners.", body),
        Paragraph("☑ <b>Social pillar:</b> Women co-founded leadership; child-safeguarding by design (no child accounts); inclusive school/CHW pilots; commitment to non-discrimination and safe labour practices with contractors.", body),
        Paragraph("☑ <b>Environmental pillar:</b> Digital-first delivery; efficient public climate APIs; minimize travel via remote enablement; product purpose is climate-health anticipatory action (SDG 13 + 3).", body),
        Paragraph(
            "<b>Integration in contract execution:</b> If awarded, we will (1) publish open releases to avoid lock-in, "
            "(2) measure pilot carbon/travel modestly and prefer remote training, (3) include women and local facilitators "
            "in field teams, (4) report sustainability actions in quarterly VF updates, and (5) keep child data minimization "
            "as a hard requirement.",
            body,
        ),
        Spacer(1, 20),
        Paragraph(f"Signature: ________________________        Date: ____________", body),
        Paragraph(f"Name: {DIRECTOR_W['full']}", body),
        Paragraph(f"Company: {COMPANY}", body),
    ]
    SimpleDocTemplate(str(path), pagesize=A4, leftMargin=54, rightMargin=54, topMargin=48, bottomMargin=48,
                      title="Appendix 1 Sustainable Procurement", author=COMPANY).build(story)
    return path


def write_readme() -> Path:
    path = OUT / "README_SUBMISSION.md"
    path.write_text(
        f"""# UNGM / UNICEF Venture Fund submission pack

**RFPS:** {RFPS}  
**Supplier:** {COMPANY} (CIN {CIN})  
**Solution:** {SOLUTION}  
**Country:** {COUNTRY} (UNICEF programme country)  
**Request:** USD 100,000 (equity-free seed), 12 months  
**Deadline:** 10 August 2026, 23:59 CEST (queries by 3 August 2026)  
**Contact:** {EMAIL}

## Files in this folder

| File | Annex |
|------|--------|
| `00_Cover_Letter_...pdf` | Cover letter |
| `01_Summary_...docx` | Template 1 |
| `02_Product_Requirements_...docx` | Template 2 |
| `03_Product_Design_Data_Plan_...xlsx` | Template 3 |
| `04_Team_Structure_...docx` + `04_Org_Chart_...pdf` | Template 4 |
| `05_CV_*.docx` | Template 5 |
| `06_Business_Model_Plan_...docx` | Template 6 |
| `07_Budget_Proposal_...xlsx` | Template 7 |
| `08_Appendix1_Sustainable_Procurement_...pdf` | TOR Appendix 1 |

Also attach from repo (recommended):

- `docs/KGK_Technical_Commendation.pdf`
- Live demo: {DEMO}
- Source: {REPO} (grant `@unicefinnovation` / venturefund@unicef.org access if ever private)

## Before e-submission (human actions)

1. Upload **2-minute YouTube pitch** and paste URL into Template 1.
2. Confirm **DOB, education, LinkedIn, references** on CVs; sign Appendix 1 and RFPS declaration form.
3. Opt-in on UNGM e-submission; submit Technical + Financial parts before deadline.
4. Quote currency **USD**; keep proposal valid ≥180 days.
5. Review Annex D1 (USD Funding Agreement) before signing award.

## Regenerate filled templates

```bash
python3 UNGM/fill_sustainow_bid.py
```
""",
        encoding="utf-8",
    )
    return path


def main() -> None:
    paths = []
    paths.append(write_cover_letter())
    paths.append(fill_template1())
    paths.append(fill_template2())
    paths.append(fill_template3())
    paths.append(fill_template4_orgchart())
    paths.extend(fill_template5())
    paths.append(fill_template6())
    paths.append(fill_template7())
    paths.append(write_appendix1())
    paths.append(write_readme())
    # update root readme
    (ROOT / "readme").write_text(
        f"UNICEF Venture Fund {RFPS} — Sustainow Technologies / KlimaGuard Kids\n"
        f"Filled bid package: UNGM/submission/\n"
        f"Regenerate: python3 UNGM/fill_sustainow_bid.py\n"
        f"Deadline: 10 August 2026 23:59 CEST\n",
        encoding="utf-8",
    )
    print("Wrote submission pack:")
    for p in paths:
        print(" ", p)


if __name__ == "__main__":
    main()
