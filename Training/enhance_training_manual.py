#!/usr/bin/env python3
"""Enhance KlimaGuard Kids Marketing Training Manual with images and context."""

from __future__ import annotations

import io
import subprocess
import textwrap
from pathlib import Path

import cairosvg
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from playwright.sync_api import sync_playwright

ROOT = Path(__file__).resolve().parents[1]
TRAINING = ROOT / "Training"
ASSETS = TRAINING / "manual_assets"
DOCX_PATH = TRAINING / "KlimaGuard_Kids_Marketing_Training_Manual.docx"
BASE_URL = "http://localhost:3000"

OCEAN = "#0ea5e9"
INK = "#0f172a"
SAFFRON = "#f97316"
LEAF = "#4ade80"
SKY = "#e0f2fe"
MUTED = "#475569"


def ensure_assets_dir() -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    return ASSETS


def export_logo() -> Path:
    out = ASSETS / "logo.png"
    cairosvg.svg2png(
        url=str(ROOT / "public" / "logo.svg"),
        write_to=str(out),
        output_width=256,
        output_height=256,
    )
    return out


def capture_screenshots() -> dict[str, Path]:
    shots: dict[str, Path] = {}
    pages = [
        ("home", "/"),
        ("dashboard", "/dashboard"),
        ("india", "/india"),
        ("chat", "/chat"),
        ("pitch", "/pitch"),
    ]
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page(viewport={"width": 1280, "height": 900})
        for name, path in pages:
            out = ASSETS / f"screenshot_{name}.png"
            page.goto(f"{BASE_URL}{path}", wait_until="networkidle")
            page.wait_for_timeout(1500)
            page.screenshot(path=str(out), full_page=True)
            shots[name] = out
        browser.close()
    return shots


def draw_agent_architecture() -> Path:
    out = ASSETS / "diagram_agent_architecture.png"
    fig, ax = plt.subplots(figsize=(12, 7))
    fig.patch.set_facecolor("white")
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    ax.axis("off")
    ax.set_title(
        "KlimaGuard Kids — Nine AI Agent Architecture",
        fontsize=16,
        fontweight="bold",
        color=INK,
        pad=16,
    )

    def box(x, y, w, h, text, color, fs=9):
        patch = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.04,rounding_size=0.15",
            facecolor=color,
            edgecolor=INK,
            linewidth=1.2,
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fs, color=INK, wrap=True)

    box(4.5, 5.8, 3, 0.7, "Open-Meteo\nClimate + Air Quality Data", SKY, 10)
    box(4.5, 4.5, 3, 0.7, "Climate Data Agent", OCEAN, 10)

    agents = [
        (0.3, 3.0, "Health Risk\nAgent"),
        (2.5, 3.0, "Nutrition &\nFood Security"),
        (4.7, 3.0, "Disease\nOutlook"),
        (6.9, 3.0, "Natural\nMedicine"),
        (9.1, 3.0, "India Regional\nContext"),
    ]
    for x, y, label in agents:
        box(x, y, 1.8, 0.9, label, "#bae6fd", 8)

    box(3.5, 1.5, 2.2, 0.8, "India Child Health\nImpact Agent (CHIS)", SAFFRON, 9)
    box(6.3, 1.5, 2.2, 0.8, "Synthesis Agent\n(Age-banded guidance)", LEAF, 9)

    box(3.8, 0.2, 4.4, 0.8, "Kids Health Chat Agent + Dashboard Outputs", "#fef3c7", 10)

    for x in [1.2, 3.4, 5.6, 7.8, 10.0]:
        ax.annotate("", xy=(x, 3.0), xytext=(6.0, 4.5), arrowprops=dict(arrowstyle="->", color=MUTED, lw=1.2))
    ax.annotate("", xy=(4.6, 1.5), xytext=(5.0, 3.0), arrowprops=dict(arrowstyle="->", color=MUTED, lw=1.2))
    ax.annotate("", xy=(7.4, 1.5), xytext=(6.0, 3.0), arrowprops=dict(arrowstyle="->", color=MUTED, lw=1.2))
    ax.annotate("", xy=(6.0, 1.0), xytext=(6.0, 1.5), arrowprops=dict(arrowstyle="->", color=MUTED, lw=1.2))

    fig.tight_layout()
    fig.savefig(out, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out


def draw_product_flow() -> Path:
    out = ASSETS / "diagram_product_flow.png"
    steps = [
        "1. Data Sources\n(Open-Meteo, IMD,\nCPCB, NFHS-5)",
        "2. Agent Pipeline\n(9 cooperating\nAI agents)",
        "3. Regional Mapping\n(65 countries /\n37 India regions)",
        "4. Impact Scoring\n(CHIS 0–100\nacross 5 dimensions)",
        "5. Dashboards\n& Chat",
        "6. Stakeholder\nAction",
    ]
    fig, ax = plt.subplots(figsize=(12, 3.2))
    fig.patch.set_facecolor("white")
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 3)
    ax.axis("off")
    ax.set_title("Product Flow — From Data to Decision Support", fontsize=14, fontweight="bold", color=INK)

    colors = [SKY, OCEAN, SAFFRON, LEAF, "#fef3c7", "#ddd6fe"]
    for i, (step, color) in enumerate(zip(steps, colors)):
        x = 0.2 + i * 1.95
        patch = FancyBboxPatch(
            (x, 0.8),
            1.7,
            1.4,
            boxstyle="round,pad=0.05,rounding_size=0.12",
            facecolor=color,
            edgecolor=INK,
            linewidth=1,
        )
        ax.add_patch(patch)
        ax.text(x + 0.85, 1.5, step, ha="center", va="center", fontsize=8, color=INK)
        if i < len(steps) - 1:
            ax.annotate(
                "",
                xy=(x + 1.75, 1.5),
                xytext=(x + 1.85, 1.5),
                arrowprops=dict(arrowstyle="->", color=INK, lw=1.5),
            )

    fig.tight_layout()
    fig.savefig(out, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out


def draw_chis_dimensions() -> Path:
    out = ASSETS / "diagram_chis_dimensions.png"
    labels = [
        "CHVI\nHeat\n(25%)",
        "CRBS\nAir Quality\n(20%)",
        "WDPI\nWaterborne\n(20%)",
        "VBDP\nVectors\n(20%)",
        "CNSI\nNutrition\n(15%)",
    ]
    sizes = [25, 20, 20, 20, 15]
    colors = [SAFFRON, OCEAN, "#38bdf8", LEAF, "#fbbf24"]

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(11, 4.5))
    fig.patch.set_facecolor("white")
    ax1.pie(sizes, labels=labels, colors=colors, autopct="%1.0f%%", startangle=90, textprops={"fontsize": 9})
    ax1.set_title("CHIS Composite Weights", fontsize=13, fontweight="bold", color=INK)

    dims = ["CHVI", "CRBS", "WDPI", "VBDP", "CNSI"]
    example = [72, 68, 55, 48, 61]
    bars = ax2.barh(dims, example, color=colors, edgecolor=INK, linewidth=0.8)
    ax2.set_xlim(0, 100)
    ax2.set_xlabel("Example Score (0–100, higher = greater burden)")
    ax2.set_title("Sample CHIS Dimension Scores — Mumbai", fontsize=13, fontweight="bold", color=INK)
    for bar, val in zip(bars, example):
        ax2.text(val + 1, bar.get_y() + bar.get_height() / 2, str(val), va="center", fontsize=9)

    fig.suptitle("Child Health Impact Score (CHIS) — Five Transparent Dimensions", fontsize=14, fontweight="bold", color=INK, y=1.02)
    fig.tight_layout()
    fig.savefig(out, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out


def draw_stakeholder_funnel() -> Path:
    out = ASSETS / "diagram_stakeholder_funnel.png"
    stages = [
        ("Awareness\nOutreach", 10.0, OCEAN),
        ("Discovery\nCall", 8.2, "#38bdf8"),
        ("Live Demo\n(15 min)", 6.4, SAFFRON),
        ("Pilot\n(10–50 sites)", 4.6, LEAF),
        ("Evidence &\nPartnership", 2.8, "#a78bfa"),
        ("Scale", 1.0, INK),
    ]
    fig, ax = plt.subplots(figsize=(10, 5))
    fig.patch.set_facecolor("white")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6.5)
    ax.axis("off")
    ax.set_title("Marketing Funnel — Right Stakeholder to Scale", fontsize=14, fontweight="bold", color=INK)

    y = 5.5
    for label, width, color in stages:
        x = (10 - width) / 2
        patch = FancyBboxPatch(
            (x, y),
            width,
            0.75,
            boxstyle="round,pad=0.03,rounding_size=0.1",
            facecolor=color,
            edgecolor=INK,
            linewidth=1,
        )
        ax.add_patch(patch)
        text_color = "white" if color == INK else INK
        ax.text(5, y + 0.375, label, ha="center", va="center", fontsize=10, fontweight="bold", color=text_color)
        y -= 0.95

    fig.tight_layout()
    fig.savefig(out, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out


def draw_india_coverage() -> Path:
    out = ASSETS / "diagram_india_coverage.png"
    tiers = {
        "Tier 1 Metros (8)": "Delhi NCR · Mumbai · Bengaluru · Hyderabad · Chennai · Kolkata · Ahmedabad · Pune",
        "Tier 2 Hubs (19)": "Lucknow · Jaipur · Kochi · Surat · Bhubaneswar · Visakhapatnam · Indore · Guwahati · Dehradun · Varanasi · Chandigarh · Nagpur · Coimbatore · Vadodara · Thiruvananthapuram · Bhopal · Amritsar · Raipur",
        "Tier 3 Centres (10)": "Patna · Ranchi · Srinagar · Kanpur · Agra · Jodhpur · Gorakhpur · Madurai · Tiruchirappalli · Jammu · Siliguri",
    }
    fig, ax = plt.subplots(figsize=(11, 5))
    fig.patch.set_facecolor("white")
    ax.axis("off")
    ax.set_title("India Regional Coverage — 37 Climate-Health Zones", fontsize=14, fontweight="bold", color=INK, pad=12)

    y = 0.85
    colors = [SAFFRON, OCEAN, LEAF]
    for (title, cities), color in zip(tiers.items(), colors):
        ax.text(0.02, y, title, fontsize=12, fontweight="bold", color=color, transform=ax.transAxes)
        wrapped = textwrap.fill(cities, width=95)
        ax.text(0.02, y - 0.08, wrapped, fontsize=9, color=INK, transform=ax.transAxes, va="top")
        y -= 0.28

    ax.text(
        0.02,
        0.05,
        "Each region includes climate zone metadata, monsoon context, and tier classification for targeted outreach.",
        fontsize=9,
        color=MUTED,
        transform=ax.transAxes,
        style="italic",
    )
    fig.tight_layout()
    fig.savefig(out, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0E, 0xA5, 0xE9)


def add_image(doc: Document, path: Path, caption: str, width: float = 6.0) -> None:
    doc.add_picture(str(path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor(0x47, 0x55, 0x69)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def enhance_document(
    logo: Path,
    screenshots: dict[str, Path],
    diagrams: dict[str, Path],
) -> None:
    doc = Document(str(DOCX_PATH))

    # Cover logo near top if possible
    doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[0].insert_paragraph_before("")
    p_logo = doc.paragraphs[0].insert_paragraph_before("")
    run = p_logo.add_run()
    run.add_picture(str(logo), width=Inches(1.2))
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insert product flow after section 5 heading
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "5. How the Product Works — Simple Flow":
            # Insert after the 8 steps — find step 8 paragraph
            pass

    # Append comprehensive new sections
    doc.add_page_break()
    add_heading(doc, "26. Platform Visual Walkthrough", 1)
    doc.add_paragraph(
        "This section provides visual context from the live KlimaGuard Kids prototype. "
        "Use these screenshots during demos and training exercises to connect marketing "
        "language with actual product surfaces."
    )

    add_heading(doc, "26.1 Home Page — Product Entry Point", 2)
    add_image(doc, screenshots["home"], "Figure 26.1: Home page with navigation to India dashboard, kids chat, global dashboard, and pitch.")
    add_bullets(doc, [
        "Marketing angle: One platform, multiple entry points for schools (chat), CSR (India impact), researchers (global dashboard).",
        "Key message: Open source, agentic AI, 37 India regions, 65 global countries.",
        "Call-to-action: Direct stakeholders to the surface most relevant to their use case.",
    ])

    add_heading(doc, "26.2 Global Climate-Health Dashboard", 2)
    add_image(doc, screenshots["dashboard"], "Figure 26.2: Global dashboard with country selector and six-agent pipeline analysis.")
    add_bullets(doc, [
        "Demonstrates live Open-Meteo integration — no API key required for demo.",
        "Six global agents produce child-ready guidance for 65 climate-vulnerable countries.",
        "For India, links to dedicated CHIS dashboard with two additional regional agents.",
    ])

    add_heading(doc, "26.3 India Child Health Impact Dashboard", 2)
    add_image(doc, screenshots["india"], "Figure 26.3: India dashboard measuring CHIS across 37 Tier 1–3 regions.")
    add_bullets(doc, [
        "Primary demo surface for Indian school groups, NGOs, CSR, and government stakeholders.",
        "Shows composite CHIS score plus five dimension breakdowns with methodology notes.",
        "References IMD, CPCB, NFHS-5, and NVBDCP — cite these for credibility with researchers.",
    ])

    add_heading(doc, "26.4 Kids Health Chat", 2)
    add_image(doc, screenshots["chat"], "Figure 26.4: Age-banded kids health chat with privacy safeguards.")
    add_bullets(doc, [
        "Age bands: 5–8, 9–12, 13–17 — answers adapt to reading level and complexity.",
        "Privacy-first: browser session only, no child accounts required for demo.",
        "Consultant scheduling available — position as escalation path, not replacement for clinical care.",
    ])

    add_heading(doc, "26.5 Product Pitch Page", 2)
    add_image(doc, screenshots["pitch"], "Figure 26.5: Stakeholder-facing product overview with SDG alignment.")
    add_bullets(doc, [
        "Use for pre-demo email attachments and association webinar follow-ups.",
        "Covers problem, nine agents, safeguarding, and open-source positioning.",
    ])

    doc.add_page_break()
    add_heading(doc, "27. Nine AI Agents — Marketing Reference Guide", 1)
    add_image(doc, diagrams["architecture"], "Figure 27.1: Nine-agent architecture showing data flow from Open-Meteo to dashboards and chat.", 6.5)
    doc.add_paragraph(
        "Marketing professionals should be able to name all nine agents and explain each in one sentence. "
        "Do not overstate AI capabilities — agents organize, correlate, and explain; they do not diagnose."
    )

    agents_table = doc.add_table(rows=10, cols=3)
    agents_table.style = "Table Grid"
    headers = ["Agent", "One-Line Marketing Explanation", "Primary Stakeholder"]
    data = [
        ["Climate Data Agent", "Fetches live 7-day weather, heat index, and air quality from Open-Meteo.", "All stakeholders"],
        ["Health Risk Agent", "Maps heat stress, respiratory, flood, and vector alerts using WHO-aligned heuristics.", "Schools, public health"],
        ["Nutrition Agent", "Assesses hydration needs, food safety, and climate-linked nutrition stress.", "NGOs, CSR (nutrition programs)"],
        ["Disease Outlook Agent", "Surfaces waterborne, vector-borne, and heat illness preparedness context.", "Public health, researchers"],
        ["Natural Medicine Agent", "Evidence-tagged supportive remedies under caregiver supervision only.", "Community health, NGOs"],
        ["Synthesis Agent", "Correlates all agent outputs into age-banded guidance (ages 5–17).", "Teachers, parents, schools"],
        ["India Regional Context Agent", "Interprets monsoon cycles and zone-specific vulnerability across 37 regions.", "Government, India CSR"],
        ["India Child Health Impact Agent", "Computes transparent CHIS score (0–100) across five dimensions.", "Researchers, CSR, policy"],
        ["Kids Health Chat Agent", "Conversational Q&A with privacy filtering and consultant scheduling.", "Schools, parents"],
    ]
    for j, h in enumerate(headers):
        agents_table.rows[0].cells[j].text = h
    for i, row in enumerate(data, start=1):
        for j, val in enumerate(row):
            agents_table.rows[i].cells[j].text = val

    doc.add_page_break()
    add_heading(doc, "28. Child Health Impact Score (CHIS) — Deep Context", 1)
    add_image(doc, diagrams["chis"], "Figure 28.1: CHIS composite weights and example dimension scores.", 6.5)
    doc.add_paragraph(
        "The CHIS (Child Health Impact Score) is a 0–100 composite index where higher scores indicate "
        "greater climate-related health burden for children in a given region. All formulas are "
        "open-source and transparent in the codebase."
    )
    add_bullets(doc, [
        "CHVI (25%): Child Heat Vulnerability Index — IMD heatwave thresholds, days above 38°C/42°C, aridity factor.",
        "CRBS (20%): Child Respiratory Burden Score — AQI exposure, PM2.5/PM10, CPCB category weighting.",
        "WDPI (20%): Waterborne Disease Pressure Index — precipitation, monsoon seasonality, flood risk.",
        "VBDP (20%): Vector-Borne Disease Pressure — temperature-humidity-rainfall vector suitability (NVBDCP-aligned).",
        "CNSI (15%): Climate Nutrition Stress Index — heat/flood/drought signals, mid-day meal disruption risk.",
    ])
    doc.add_paragraph(
        "Approved marketing language: 'CHIS helps prioritize regional preparedness discussions.' "
        "Avoid: 'CHIS predicts which children will get sick.' Scores reflect regional context, not individual outcomes."
    )

    add_heading(doc, "29. India Regional Coverage Map", 1)
    add_image(doc, diagrams["india"], "Figure 29.1: 37 Indian regions organized by tier for targeted outreach.", 6.5)
    doc.add_paragraph(
        "When prospecting, match tier to stakeholder size: Tier 1 for large school networks and metro CSR; "
        "Tier 2 for state-level NGOs and regional government; Tier 3 for community programs and district pilots."
    )

    add_heading(doc, "30. 15-Minute Marketing Demo Script", 1)
    add_image(doc, diagrams["flow"], "Figure 30.1: End-to-end product flow from data sources to stakeholder action.", 6.5)
    demo_steps = [
        ("Minutes 0–3: Discovery", "Ask: Who are your children/students? Which geography? What climate-health concerns matter most? Do not demo until you understand their context."),
        ("Minutes 3–7: India Dashboard", "Select their nearest region. Run analysis. Walk through CHIS score and top two dimensions. Explain methodology transparency."),
        ("Minutes 7–10: Age-Banded Guidance", "Show synthesis output for their target age band. Emphasize preparedness, not prediction."),
        ("Minutes 10–12: Kids Chat (optional)", "Demo age-appropriate Q&A. Highlight privacy and safeguarding."),
        ("Minutes 12–15: Next Steps", "Propose defined pilot: 10-school cohort, 90-day timeline, agreed outputs (monthly CHIS reports, teacher briefing cards)."),
    ]
    for title, detail in demo_steps:
        p = doc.add_paragraph()
        p.add_run(title + ": ").bold = True
        p.add_run(detail)

    add_heading(doc, "31. Competitive Differentiation", 1)
    diff_table = doc.add_table(rows=6, cols=3)
    diff_table.style = "Table Grid"
    diff_data = [
        ["Capability", "Weather App", "KlimaGuard Kids"],
        ["Child-centric framing", "No", "Yes — ages 5–17 guidance"],
        ["Multi-domain correlation", "No", "Yes — health, nutrition, disease, climate"],
        ["Regional impact scoring", "No", "Yes — CHIS across 37 India regions"],
        ["Transparent methodology", "N/A", "Yes — open-source MIT, formulas published"],
        ["Age-banded chat", "No", "Yes — privacy-safe kids health Q&A"],
    ]
    for i, row in enumerate(diff_data):
        for j, val in enumerate(row):
            diff_table.rows[i].cells[j].text = val

    add_heading(doc, "32. SDG & CSR Alignment Talking Points", 1)
    add_bullets(doc, [
        "SDG 3 (Good Health): Anticipatory action for climate-sensitive child health risks.",
        "SDG 13 (Climate Action): Child-centred early warning intelligence layer.",
        "SDG 2 (Zero Hunger): Nutrition and food security inference during climate shocks.",
        "SDG 4 (Quality Education): School-ready heat, flood, and air quality guidance for administrators.",
        "CSR framing: Fund a measurable 10/25/50-school pilot with monthly CHIS reporting and teacher preparedness cards.",
    ])

    add_heading(doc, "33. Open Source & Partnership Positioning", 1)
    add_bullets(doc, [
        "MIT License — free to use, modify, and distribute; lowers adoption friction for NGOs and researchers.",
        "Transparent formulas in open codebase — invite technical review rather than asking stakeholders to trust black boxes.",
        "Extensible via API: /api/analyze, /api/chat, /api/india/regions — supports integration with existing school or health systems.",
        "Community contributions welcome: new regions, agents, language packs, offline modes.",
        "Positioning: 'Intelligence layer you can inspect, adapt, and co-develop' — not a locked proprietary SaaS.",
    ])

    add_heading(doc, "34. Marketing Funnel Visual Reference", 1)
    add_image(doc, diagrams["funnel"], "Figure 34.1: Standard funnel from awareness outreach to scale partnership.", 5.0)
    doc.add_paragraph(
        "Align every touchpoint to the funnel stage. Early emails should not ask for enterprise contracts; "
        "they should earn a 15-minute discovery call. Pilots must have defined outputs before scale conversations."
    )

    add_heading(doc, "35. Quick Reference Card — Trainer Handout", 1)
    ref_table = doc.add_table(rows=8, cols=2)
    ref_table.style = "Table Grid"
    ref_items = [
        ["Product type", "Child-centric climate-health intelligence & decision support"],
        ["Not a", "Medical diagnostic, treatment system, or individual disease predictor"],
        ["Live demo URL paths", "/ · /india · /dashboard · /chat · /pitch"],
        ["India coverage", "37 regions (8 Tier 1 · 19 Tier 2 · 10 Tier 3)"],
        ["Global coverage", "65 climate-vulnerable countries"],
        ["AI agents", "9 cooperating agents with transparent orchestration"],
        ["Key score", "CHIS 0–100 (higher = greater regional burden)"],
        ["License", "MIT Open Source"],
    ]
    for i, (k, v) in enumerate(ref_items):
        ref_table.rows[i].cells[0].text = k
        ref_table.rows[i].cells[1].text = v

    doc.save(str(DOCX_PATH))
    print(f"Enhanced manual saved to {DOCX_PATH}")


def main() -> None:
    ensure_assets_dir()
    print("Exporting logo...")
    logo = export_logo()
    print("Capturing screenshots...")
    screenshots = capture_screenshots()
    print("Generating diagrams...")
    diagrams = {
        "architecture": draw_agent_architecture(),
        "flow": draw_product_flow(),
        "chis": draw_chis_dimensions(),
        "funnel": draw_stakeholder_funnel(),
        "india": draw_india_coverage(),
    }
    print("Enhancing document...")
    enhance_document(logo, screenshots, diagrams)
    print("Done.")


if __name__ == "__main__":
    main()
