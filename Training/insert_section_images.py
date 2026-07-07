#!/usr/bin/env python3
"""Insert contextual images into existing training manual sections."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

TRAINING = Path(__file__).resolve().parent
DOCX_PATH = TRAINING / "KlimaGuard_Kids_Marketing_Training_Manual.docx"
ASSETS = TRAINING / "manual_assets"


def insert_after_heading(doc: Document, heading_text: str, image_path: Path, caption: str, width: float = 5.5) -> bool:
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == heading_text:
            new_p = para._element
            # Insert picture paragraph after heading
            from docx.oxml.ns import qn
            from docx.text.paragraph import Paragraph

            body = doc.element.body
            idx = list(body).index(new_p) + 1

            pic_p = doc.add_paragraph()
            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pic_p.add_run()
            run.add_picture(str(image_path), width=Inches(width))
            body.insert(idx, pic_p._element)
            idx += 1

            cap_p = doc.add_paragraph(caption)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cap_p.runs:
                cap_p.runs[0].italic = True
                cap_p.runs[0].font.size = Pt(9)
                cap_p.runs[0].font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            body.insert(idx, cap_p._element)
            return True
    return False


def insert_context_paragraph(doc: Document, after_heading: str, text: str) -> bool:
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == after_heading:
            from docx.text.paragraph import Paragraph
            body = doc.element.body
            idx = list(body).index(para._element) + 1
            ctx = doc.add_paragraph(text)
            body.insert(idx, ctx._element)
            return True
    return False


def main() -> None:
    doc = Document(str(DOCX_PATH))

    insertions = [
        (
            "2. Product Overview",
            ASSETS / "diagram_product_flow.png",
            "Figure 2.1: High-level product flow — from live climate data to stakeholder decision support.",
        ),
        (
            "5. How the Product Works — Simple Flow",
            ASSETS / "diagram_agent_architecture.png",
            "Figure 5.1: Nine-agent pipeline architecture — reference during the eight-step flow explanation.",
        ),
        (
            "7. AI, Dashboards and Impact Scoring",
            ASSETS / "diagram_chis_dimensions.png",
            "Figure 7.1: CHIS composite score breakdown across five transparent dimensions.",
        ),
        (
            "8. Target Customer Profiles",
            ASSETS / "diagram_stakeholder_funnel.png",
            "Figure 8.1: Marketing funnel — align outreach stage to stakeholder readiness.",
        ),
        (
            "25. Final Message for the Marketing Team",
            ASSETS / "diagram_india_coverage.png",
            "Figure 25.1: India regional coverage — 37 Tier 1–3 cities for localized pitches.",
        ),
    ]

    for heading, img, caption in insertions:
        ok = insert_after_heading(doc, heading, img, caption)
        print(f"{'OK' if ok else 'MISS'}: {heading}")

    insert_context_paragraph(
        doc,
        "2. Product Overview",
        "Platform context: KlimaGuard Kids is an open-source (MIT) Next.js web application with live "
        "Open-Meteo integration, nine cooperating AI agents, 65 global countries, and 37 India regions. "
        "The live prototype includes home, global dashboard, India CHIS dashboard, kids health chat, "
        "and stakeholder pitch pages — all accessible without API keys for demonstration purposes.",
    )

    insert_context_paragraph(
        doc,
        "6. Climate-Health Fundamentals for Marketing Professionals",
        "India-specific context: The India Child Health Impact Agent uses IMD heatwave thresholds, "
        "CPCB air quality categories, NFHS-5 child health baselines, and NVBDCP vector seasonality "
        "models. When discussing these concepts with Indian stakeholders, reference these authoritative "
        "sources to build credibility — but always clarify that the platform synthesizes public data "
        "into decision support, not clinical diagnoses.",
    )

    insert_context_paragraph(
        doc,
        "13. Email Marketing Training",
        "Visual demo tip: Include a screenshot of the relevant dashboard (India CHIS for Indian schools, "
        "global dashboard for international NGOs) as an inline image or link to /india or /dashboard. "
        "Personalized visuals increase demo booking rates compared to text-only outreach.",
    )

    doc.save(str(DOCX_PATH))
    print(f"Updated {DOCX_PATH}")


if __name__ == "__main__":
    main()
